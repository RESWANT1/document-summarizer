"""
extractors/docx_extractor.py
Extracts plain text from .docx Word documents using python-docx.
Preserves paragraph order, joins run-on lines within a paragraph, and
includes table cell content as coherent rows (not isolated cells).
"""

import logging
from docx import Document

logger = logging.getLogger(__name__)


def extract_docx(filepath: str) -> str:
    """Return full text from a .docx file (paragraphs + tables).

    Improvements over naïve extraction:
    - Consecutive non-empty paragraphs that look like they belong to the same
      logical sentence (no terminal punctuation, short next line) are joined
      with a space so the downstream sentence segmenter sees them as one unit.
    - Table rows are joined on a single line separated by ' | ' so the
      sentence tokeniser treats an entire row as one record instead of
      producing dangling single-word fragments.
    """
    doc = Document(filepath)
    parts: list[str] = []

    # ── Paragraphs ────────────────────────────────────────────────────────────
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    i = 0
    while i < len(para_texts):
        current = para_texts[i]
        # If the line has no terminal punctuation and the next line is short
        # (likely a continuation — e.g. "I will be based in the U." → "S.")
        # merge them together.
        while (
            i + 1 < len(para_texts)
            and current
            and current[-1] not in ".!?:;"
            and len(para_texts[i + 1].split()) <= 6
        ):
            i += 1
            current = current + " " + para_texts[i]
        parts.append(current)
        i += 1

    # ── Table cells ──────────────────────────────────────────────────────────
    # Join cells within a row with ' | ' so BART sees "Role | Apr 2022 – Present"
    # rather than three detached fragments.
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            seen = set()       # docx can return merged cells twice
            for cell in row.cells:
                txt = cell.text.strip()
                if txt and txt not in seen:
                    seen.add(txt)
                    row_cells.append(txt)
            if row_cells:
                parts.append(" | ".join(row_cells))

    full_text = "\n".join(parts)
    logger.info("DOCX extraction complete: %d characters.", len(full_text))
    return full_text